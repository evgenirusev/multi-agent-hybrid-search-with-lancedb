#!/usr/bin/env python3

import os
import tempfile
import logging
import time
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field, field_validator
from src.agents.crews.legal_support_agents.legal_support_agents import LegalSupportAgents
from src.agents.rag import document_store, initialize_document_store
import uvicorn
from dotenv import load_dotenv
from docx import Document
from contextlib import asynccontextmanager
import re

# Configure minimal logging
logging.basicConfig(
    level=logging.WARNING,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("api")

# Load environment variables from .env file
load_dotenv()

# Initialize FastAPI app with lifespan
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup: Initialize document store
    try:
        await initialize_document_store()
    except Exception as e:
        logger.error(f"Error initializing document store: {e}")
        logger.warning("API will start, but document storage functionality may not work properly")
    
    yield  # This is where FastAPI serves the application

# Initialize FastAPI app
app = FastAPI(
    title="Agentic API", 
    description="API for AI services",
    lifespan=lifespan
)

# Simplified request logging middleware
@app.middleware("http")
async def log_requests(request: Request, call_next):
    start_time = time.time()
    try:
        response = await call_next(request)
        process_time = time.time() - start_time
        response.headers["X-Process-Time"] = str(process_time)
        return response
    except Exception as e:
        logger.error(f"Request failed: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"detail": "Internal Server Error", "error": str(e)}
        )

# Add CORS middleware to allow cross-origin requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

class QueryRequest(BaseModel):
    Query: str = Field(..., min_length=1, max_length=2000, alias="query")  # Add length validation

    class Config:
        populate_by_name = True

class QueryResponse(BaseModel):
    result: str

class SearchRequest(BaseModel):
    query: str = Field(..., min_length=1, max_length=2000)  # Add length validation
    limit: int = Field(5, ge=1, le=20)  # Add range validation
    
    @field_validator('query')
    @classmethod
    def validate_query(cls, v):
        # Basic security validation
        if re.search(r'[<>{}]', v):  # Detect potential HTML/script injection
            raise ValueError('Query contains invalid characters')
        return v

@app.post("/query", response_model=QueryResponse)
async def process_query(request: QueryRequest):
    try:
        legal_crew = LegalSupportAgents(debug_enabled=False)

        # Use the fully async process_query method
        result = await legal_crew.process_query(request.Query)

        return QueryResponse(result=result)
    except Exception as crew_error:
        logger.error(f"Error in LegalSupportCrew: {str(crew_error)}")
        fallback_response = "I apologize, but I'm currently experiencing technical difficulties. Please try again later."
        return QueryResponse(result=fallback_response)

@app.post("/vectorize-document")
async def vectorize_document(
    file: UploadFile = File(...), 
    document_name: str = Form(None)
):
    # Check if the file is a Word document
    if not file.filename.endswith('.docx'):
        logger.warning(f"Invalid file type: {file.filename}")
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        # Use filename as document_name if not provided
        if not document_name:
            document_name = file.filename
        else:
            # Basic validation for document_name
            if len(document_name) > 200:
                raise HTTPException(status_code=400, detail="Document name too long")
            
        # Create a temporary file with secure random name
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            # Read the uploaded file
            content = await file.read()
            
            # Check file size limit (10MB)
            if len(content) > 10 * 1024 * 1024:  # 10MB in bytes
                raise HTTPException(status_code=400, detail="File too large")
                
            # Write to the temporary file
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text from the document
            document_text = extract_text_from_docx(temp_file_path)
            text_preview = document_text[:100] + "..." if len(document_text) > 100 else document_text
            
            # Add the document to the vector store
            result = await document_store.add_document(document_text, document_name)
            
            return JSONResponse(content={
                "filename": file.filename,
                "document_name": document_name,
                "document_id": result["document_id"],
                "document_text": text_preview,
                "chunks_added": result["chunks_added"]
            })
        finally:
            # Always clean up the temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
    
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        # Don't expose detailed error information
        raise HTTPException(status_code=500, detail="Error processing document")

@app.post("/embeddings")
async def get_document_embeddings(request: SearchRequest):
    """
    Search for documents or sections matching the query.
    """
    try:
        # Search the vector store
        results = await document_store.search(
            request.query, 
            limit=request.limit
        )
        
        # Format results for response
        formatted_results = []
        for result in results:
            formatted_results.append({
                "text": result["text"],
                "document_id": result["document_id"],
                "document_name": result["document_name"],
                "section": result["section"] if "section" in result and result["section"] else None,
                "score": float(result["_relevance_score"]) if "_relevance_score" in result else 0.0
            })
        
        return JSONResponse(content={
            "query": request.query,
            "results": formatted_results
        })
        
    except Exception as e:
        logger.error(f"Error in get_document_embeddings: {e}")
        raise HTTPException(status_code=500, detail="Error searching documents")

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only count non-empty paragraphs
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Only count non-empty cells
                        full_text.append(cell.text)
        
        combined_text = "\n".join(full_text)
        return combined_text
    
    except Exception as e:
        raise Exception(f"Failed to extract text from document: {str(e)}")

@app.delete("/document/{document_id}")
async def delete_document(document_id: str):
    try:
        # Delete the document from vector store
        result = await document_store.delete_document(document_id)
        
        return JSONResponse(content={
            "document_id": document_id,
            "chunks_deleted": result["chunks_deleted"],
            "success": True
        })
        
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail="Error deleting document")

@app.get("/documents")
async def get_all_documents():
    try:
        # Get all documents from vector store
        documents = await document_store.get_all_documents()
        
        return JSONResponse(content={
            "document_count": len(documents),
            "documents": documents
        })
        
    except Exception as e:
        logger.error(f"Error retrieving documents: {e}")
        raise HTTPException(status_code=500, detail="Error retrieving documents")

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True) 